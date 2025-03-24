from flask import Flask, render_template_string, request, send_file
from PIL import Image
from docx import Document
import os

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
    <title>Image to PDF/Word Converter</title>
    <style>
        body {
            font-family: 'Poppins', sans-serif;
            text-align: center;
            margin: 50px;
            background: url('https://source.unsplash.com/random/1600x900') no-repeat center center fixed;
            background-size: cover;
            color: white;
            animation: fadeIn 2s ease-in-out;
        }
        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }
        .container {
            max-width: 450px;
            margin: auto;
            padding: 20px;
            background: rgba(0, 0, 0, 0.8);
            border-radius: 15px;
            box-shadow: 0px 4px 15px rgba(255, 255, 255, 0.2);
            animation: slideIn 1.5s ease-in-out;
        }
        @keyframes slideIn {
            from { transform: translateY(-50px); opacity: 0; }
            to { transform: translateY(0); opacity: 1; }
        }
        input[type="file"] {
            display: none;
        }
        .custom-file-upload {
            display: inline-block;
            padding: 12px;
            background: rgba(255, 255, 255, 0.3);
            color: white;
            border-radius: 8px;
            cursor: pointer;
            font-size: 16px;
            margin: 10px 0;
            transition: all 0.3s;
        }
        .custom-file-upload:hover {
            background: rgba(255, 255, 255, 0.5);
        }
        select, button {
            width: 100%;
            margin: 10px 0;
            padding: 12px;
            border: none;
            border-radius: 8px;
            font-size: 16px;
        }
        select {
            background: rgba(255, 255, 255, 0.3);
            color: white;
            outline: none;
        }
        button {
            background: linear-gradient(45deg, #ff9800, #ff5722);
            color: white;
            font-weight: bold;
            cursor: pointer;
            transition: all 0.3s;
            box-shadow: 0px 4px 10px rgba(255, 87, 34, 0.3);
        }
        button:hover {
            background: linear-gradient(45deg, #e68900, #e64a19);
            transform: scale(1.05);
        }
    </style>
</head>
<body>
    <h2>Image to PDF/Word Converter</h2>
    <div class=\"container\">
        <form action=\"/convert\" method=\"post\" enctype=\"multipart/form-data\">
            <label for=\"file-upload\" class=\"custom-file-upload\">Choose File</label>
            <input id=\"file-upload\" type=\"file\" name=\"image\" required>
            <select name=\"format\">
                <option value=\"pdf\">PDF</option>
                <option value=\"docx\">Word (DOCX)</option>
            </select>
            <button type=\"submit\">Convert</button>
        </form>
    </div>
</body>
</html>
"""


@app.route('/')
def home():
    return render_template_string(HTML_TEMPLATE)


@app.route('/convert', methods=['POST'])
def convert():
    if 'image' not in request.files:
        return "No file part"

    file = request.files['image']
    if file.filename == '':
        return "No selected file"

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    output_format = request.form.get('format')
    output_path = os.path.join(OUTPUT_FOLDER, f"output.{output_format}")

    if output_format == 'pdf':
        image = Image.open(file_path)
        image = image.convert('RGB')
        image.save(output_path, "PDF")
    elif output_format == 'docx':
        doc = Document()
        doc.add_paragraph("Image to Word Conversion")
        doc.add_picture(file_path)
        doc.save(output_path)
    else:
        return "Invalid format"

    return send_file(output_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
